import zipfile
import rarfile
import py7zr
import tarfile
import gzip
import bz2
import os
import json
import csv
import pandas as pd
from pathlib import Path
import tempfile
import shutil
from typing import List, Dict, Any, Optional
import mimetypes
from PIL import Image
import PyPDF2
import docx
from openpyxl import load_workbook

class AdvancedFileProcessor:
    """معالج ملفات متقدم لنظام 3RBAI"""
    
    def __init__(self):
        self.supported_archives = ['.zip', '.rar', '.7z', '.tar', '.gz', '.bz2']
        self.supported_images = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp', '.tiff']
        self.supported_documents = ['.pdf', '.docx', '.doc', '.xlsx', '.xls', '.pptx', '.ppt']
        self.supported_text = ['.txt', '.md', '.csv', '.json', '.xml', '.html', '.css', '.js', '.ts', '.py']
        
        print("🔧 تم تهيئة معالج الملفات المتقدم لـ 3RBAI")
    
    def process_file(self, file_path: str, output_dir: str = None) -> Dict[str, Any]:
        """معالجة ملف شاملة"""
        try:
            if output_dir is None:
                output_dir = tempfile.mkdtemp()
            
            file_path = Path(file_path)
            file_extension = file_path.suffix.lower()
            
            print(f"🔍 بدء معالجة الملف: {file_path.name}")
            
            result = {
                'file_name': file_path.name,
                'file_size': file_path.stat().st_size,
                'file_type': file_extension,
                'mime_type': mimetypes.guess_type(str(file_path))[0],
                'status': 'processing',
                'extracted_files': [],
                'analysis': {},
                'content': '',
                'metadata': {}
            }
            
            # معالجة الملفات المضغوطة
            if file_extension in self.supported_archives:
                result = self.extract_archive(file_path, output_dir, result)
            
            # معالجة الصور
            elif file_extension in self.supported_images:
                result = self.process_image(file_path, result)
            
            # معالجة المستندات
            elif file_extension in self.supported_documents:
                result = self.process_document(file_path, result)
            
            # معالجة الملفات النصية
            elif file_extension in self.supported_text:
                result = self.process_text_file(file_path, result)
            
            # معالجة عامة للملفات الأخرى
            else:
                result = self.process_generic_file(file_path, result)
            
            result['status'] = 'completed'
            print(f"✅ تم معالجة الملف بنجاح: {file_path.name}")
            
            return result
            
        except Exception as e:
            print(f"❌ خطأ في معالجة الملف {file_path}: {str(e)}")
            result['status'] = 'error'
            result['error'] = str(e)
            return result
    
    def extract_archive(self, file_path: Path, output_dir: str, result: Dict) -> Dict:
        """استخراج الملفات المضغوطة"""
        try:
            extract_dir = Path(output_dir) / f"extracted_{file_path.stem}"
            extract_dir.mkdir(exist_ok=True)
            
            file_extension = file_path.suffix.lower()
            
            print(f"📦 فك ضغط الملف: {file_path.name}")
            
            if file_extension == '.zip':
                with zipfile.ZipFile(file_path, 'r') as zip_ref:
                    zip_ref.extractall(extract_dir)
                    extracted_files = zip_ref.namelist()
            
            elif file_extension == '.rar':
                with rarfile.RarFile(file_path, 'r') as rar_ref:
                    rar_ref.extractall(extract_dir)
                    extracted_files = rar_ref.namelist()
            
            elif file_extension == '.7z':
                with py7zr.SevenZipFile(file_path, 'r') as seven_ref:
                    seven_ref.extractall(extract_dir)
                    extracted_files = seven_ref.getnames()
            
            elif file_extension in ['.tar', '.tar.gz', '.tgz']:
                with tarfile.open(file_path, 'r:*') as tar_ref:
                    tar_ref.extractall(extract_dir)
                    extracted_files = tar_ref.getnames()
            
            elif file_extension == '.gz':
                with gzip.open(file_path, 'rb') as gz_ref:
                    output_file = extract_dir / file_path.stem
                    with open(output_file, 'wb') as out_file:
                        shutil.copyfileobj(gz_ref, out_file)
                    extracted_files = [file_path.stem]
            
            elif file_extension == '.bz2':
                with bz2.open(file_path, 'rb') as bz2_ref:
                    output_file = extract_dir / file_path.stem
                    with open(output_file, 'wb') as out_file:
                        shutil.copyfileobj(bz2_ref, out_file)
                    extracted_files = [file_path.stem]
            
            # تحليل الملفات المستخرجة
            extracted_info = []
            total_size = 0
            file_types = {}
            
            for extracted_file in extracted_files:
                if extracted_file.endswith('/'):  # مجلد
                    continue
                
                extracted_path = extract_dir / extracted_file
                if extracted_path.exists():
                    file_size = extracted_path.stat().st_size
                    file_type = extracted_path.suffix.lower()
                    
                    extracted_info.append({
                        'name': extracted_file,
                        'size': file_size,
                        'type': file_type,
                        'path': str(extracted_path)
                    })
                    
                    total_size += file_size
                    file_types[file_type] = file_types.get(file_type, 0) + 1
            
            result['extracted_files'] = extracted_info
            result['analysis'] = {
                'total_extracted': len(extracted_info),
                'total_size': total_size,
                'file_types': file_types,
                'extract_directory': str(extract_dir)
            }
            
            # معالجة الملفات المستخرجة المهمة
            important_content = []
            for file_info in extracted_info[:10]:  # أول 10 ملفات
                try:
                    file_path = Path(file_info['path'])
                    if file_path.suffix.lower() in self.supported_text:
                        content = self.extract_text_content(file_path)
                        if content:
                            important_content.append(f"=== {file_info['name']} ===\n{content[:500]}...")
                except:
                    continue
            
            result['content'] = '\n\n'.join(important_content)
            
            print(f"✅ تم استخراج {len(extracted_info)} ملف من الأرشيف")
            
        except Exception as e:
            print(f"❌ خطأ في فك الضغط: {str(e)}")
            result['analysis'] = {'error': str(e)}
        
        return result
    
    def process_image(self, file_path: Path, result: Dict) -> Dict:
        """معالجة الصور"""
        try:
            print(f"🖼️ تحليل الصورة: {file_path.name}")
            
            with Image.open(file_path) as img:
                result['metadata'] = {
                    'width': img.width,
                    'height': img.height,
                    'mode': img.mode,
                    'format': img.format,
                    'has_transparency': img.mode in ('RGBA', 'LA') or 'transparency' in img.info
                }
                
                # تحليل الألوان الأساسية
                colors = img.getcolors(maxcolors=256*256*256)
                if colors:
                    dominant_colors = sorted(colors, key=lambda x: x[0], reverse=True)[:5]
                    result['metadata']['dominant_colors'] = [color[1] for color in dominant_colors]
                
                result['analysis'] = {
                    'resolution': f"{img.width}x{img.height}",
                    'aspect_ratio': round(img.width / img.height, 2),
                    'total_pixels': img.width * img.height,
                    'color_depth': len(img.getbands()),
                    'estimated_quality': 'high' if img.width > 1920 else 'medium' if img.width > 720 else 'low'
                }
                
                result['content'] = f"صورة بدقة {img.width}x{img.height} بكسل، تنسيق {img.format}"
                
        except Exception as e:
            print(f"❌ خطأ في معالجة الصورة: {str(e)}")
            result['analysis'] = {'error': str(e)}
        
        return result
    
    def process_document(self, file_path: Path, result: Dict) -> Dict:
        """معالجة المستندات"""
        try:
            print(f"📄 تحليل المستند: {file_path.name}")
            
            file_extension = file_path.suffix.lower()
            content = ""
            metadata = {}
            
            if file_extension == '.pdf':
                content, metadata = self.extract_pdf_content(file_path)
            elif file_extension in ['.docx', '.doc']:
                content, metadata = self.extract_word_content(file_path)
            elif file_extension in ['.xlsx', '.xls']:
                content, metadata = self.extract_excel_content(file_path)
            
            result['content'] = content
            result['metadata'] = metadata
            result['analysis'] = {
                'word_count': len(content.split()) if content else 0,
                'char_count': len(content) if content else 0,
                'estimated_reading_time': max(1, len(content.split()) // 200) if content else 0
            }
            
        except Exception as e:
            print(f"❌ خطأ في معالجة المستند: {str(e)}")
            result['analysis'] = {'error': str(e)}
        
        return result
    
    def process_text_file(self, file_path: Path, result: Dict) -> Dict:
        """معالجة الملفات النصية"""
        try:
            print(f"📝 تحليل الملف النصي: {file_path.name}")
            
            content = self.extract_text_content(file_path)
            
            result['content'] = content
            result['analysis'] = {
                'line_count': len(content.splitlines()) if content else 0,
                'word_count': len(content.split()) if content else 0,
                'char_count': len(content) if content else 0,
                'encoding': 'utf-8'  # افتراضي
            }
            
            # تحليل خاص للملفات المنظمة
            file_extension = file_path.suffix.lower()
            if file_extension == '.csv':
                result['analysis'].update(self.analyze_csv(file_path))
            elif file_extension == '.json':
                result['analysis'].update(self.analyze_json(file_path))
            
        except Exception as e:
            print(f"❌ خطأ في معالجة الملف النصي: {str(e)}")
            result['analysis'] = {'error': str(e)}
        
        return result
    
    def process_generic_file(self, file_path: Path, result: Dict) -> Dict:
        """معالجة عامة للملفات"""
        try:
            print(f"📎 معالجة عامة للملف: {file_path.name}")
            
            result['analysis'] = {
                'file_type': 'unknown',
                'can_process': False,
                'suggestions': ['يمكن محاولة تحويل الملف لتنسيق مدعوم', 'فحص محتوى الملف يدوياً']
            }
            
            # محاولة قراءة بداية الملف لتحديد النوع
            try:
                with open(file_path, 'rb') as f:
                    header = f.read(1024)
                    result['metadata'] = {
                        'header_preview': header[:100].hex(),
                        'is_binary': not all(32 <= byte <= 126 or byte in [9, 10, 13] for byte in header[:100])
                    }
            except:
                pass
                
        except Exception as e:
            print(f"❌ خطأ في المعالجة العامة: {str(e)}")
            result['analysis'] = {'error': str(e)}
        
        return result
    
    def extract_text_content(self, file_path: Path) -> str:
        """استخراج المحتوى النصي"""
        try:
            encodings = ['utf-8', 'utf-16', 'cp1256', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            
            # إذا فشلت جميع الترميزات، اقرأ كـ binary
            with open(file_path, 'rb') as f:
                content = f.read()
                return content.decode('utf-8', errors='ignore')
                
        except Exception as e:
            print(f"❌ خطأ في قراءة الملف النصي: {str(e)}")
            return ""
    
    def extract_pdf_content(self, file_path: Path) -> tuple:
        """استخراج محتوى PDF"""
        try:
            content = ""
            metadata = {}
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                metadata = {
                    'pages': len(pdf_reader.pages),
                    'title': pdf_reader.metadata.get('/Title', '') if pdf_reader.metadata else '',
                    'author': pdf_reader.metadata.get('/Author', '') if pdf_reader.metadata else '',
                    'creator': pdf_reader.metadata.get('/Creator', '') if pdf_reader.metadata else ''
                }
                
                for page in pdf_reader.pages:
                    content += page.extract_text() + "\n"
            
            return content, metadata
            
        except Exception as e:
            print(f"❌ خطأ في قراءة PDF: {str(e)}")
            return "", {'error': str(e)}
    
    def extract_word_content(self, file_path: Path) -> tuple:
        """استخراج محتوى Word"""
        try:
            doc = docx.Document(file_path)
            
            content = ""
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            metadata = {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'sections': len(doc.sections)
            }
            
            return content, metadata
            
        except Exception as e:
            print(f"❌ خطأ في قراءة Word: {str(e)}")
            return "", {'error': str(e)}
    
    def extract_excel_content(self, file_path: Path) -> tuple:
        """استخراج محتوى Excel"""
        try:
            workbook = load_workbook(file_path, read_only=True)
            
            content = ""
            metadata = {
                'sheets': workbook.sheetnames,
                'total_sheets': len(workbook.sheetnames)
            }
            
            for sheet_name in workbook.sheetnames[:3]:  # أول 3 أوراق فقط
                sheet = workbook[sheet_name]
                content += f"=== {sheet_name} ===\n"
                
                for row in sheet.iter_rows(max_row=100, values_only=True):  # أول 100 صف
                    row_data = [str(cell) if cell is not None else "" for cell in row]
                    content += "\t".join(row_data) + "\n"
                
                content += "\n"
            
            return content, metadata
            
        except Exception as e:
            print(f"❌ خطأ في قراءة Excel: {str(e)}")
            return "", {'error': str(e)}
    
    def analyze_csv(self, file_path: Path) -> Dict:
        """تحليل ملف CSV"""
        try:
            df = pd.read_csv(file_path, nrows=1000)  # أول 1000 صف
            
            return {
                'csv_rows': len(df),
                'csv_columns': len(df.columns),
                'csv_column_names': list(df.columns),
                'csv_data_types': df.dtypes.to_dict(),
                'csv_null_values': df.isnull().sum().to_dict()
            }
            
        except Exception as e:
            return {'csv_error': str(e)}
    
    def analyze_json(self, file_path: Path) -> Dict:
        """تحليل ملف JSON"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            def count_items(obj):
                if isinstance(obj, dict):
                    return sum(count_items(v) for v in obj.values()) + len(obj)
                elif isinstance(obj, list):
                    return sum(count_items(item) for item in obj) + len(obj)
                else:
                    return 1
            
            return {
                'json_type': type(data).__name__,
                'json_keys': list(data.keys()) if isinstance(data, dict) else [],
                'json_total_items': count_items(data),
                'json_structure_depth': self.get_json_depth(data)
            }
            
        except Exception as e:
            return {'json_error': str(e)}
    
    def get_json_depth(self, obj, depth=0):
        """حساب عمق هيكل JSON"""
        if isinstance(obj, dict):
            return max([self.get_json_depth(v, depth + 1) for v in obj.values()], default=depth)
        elif isinstance(obj, list):
            return max([self.get_json_depth(item, depth + 1) for item in obj], default=depth)
        else:
            return depth

# تشغيل المعالج
if __name__ == "__main__":
    processor = AdvancedFileProcessor()
    
    # مثال على الاستخدام
    print("🚀 معالج الملفات المتقدم لـ 3RBAI جاهز للعمل!")
    print("📁 الأنواع المدعومة:")
    print(f"   • ملفات مضغوطة: {', '.join(processor.supported_archives)}")
    print(f"   • صور: {', '.join(processor.supported_images)}")
    print(f"   • مستندات: {', '.join(processor.supported_documents)}")
    print(f"   • ملفات نصية: {', '.join(processor.supported_text)}")
